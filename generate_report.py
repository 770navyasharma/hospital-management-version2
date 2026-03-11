import sys
import subprocess
import os

try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_report():
    doc = Document()

    # Add Title
    title = doc.add_heading('Hospital Management System - Project Report', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add Student Details
    doc.add_heading('1. Student Details', level=1)
    doc.add_paragraph('Name: Navya Sharma')
    doc.add_paragraph('Roll Number: 23f1000795')

    # Add Project Details
    doc.add_heading('2. Project Details', level=1)
    doc.add_heading('Question Statement', level=2)
    doc.add_paragraph("The objective of this project is to develop a comprehensive Hospital Management System (HMS) with specialized modular portals for Patients, Doctors, and Administrators. The system streamlines scheduling, management, and secure medical record-keeping while providing role-based interactions.")

    doc.add_heading('Approach to the Problem Statement', level=2)
    doc.add_paragraph("The development approach was modular, focusing on role-based interfaces:")
    doc.add_paragraph("1. Requirements Gathering: Extracted the core personas (Admin, Doctor, Patient) and modelled their interactions around entities like Departments, Appointments, and Treatments.")
    doc.add_paragraph("2. Database Design: Constructed an interconnected relational schema using SQLAlchemy with defined foreign-key relationships to trace patient flow seamlessly.")
    doc.add_paragraph("3. Backend Strategy: Developed a robust backend using Flask and Flask-Security to handle authentication, RBAC, and RESTful API endpoints securely.")
    doc.add_paragraph("4. User Interface Architecture: Built responsive front-end dashboards using Vue.js for reactivity, Bootstrap 5 for layout, and Chart.js to map large-scale administration analytics.")
    doc.add_paragraph("5. Unique Features Implementation: Developed features prioritizing clinical usefulness, including automated PDF clinical report generation, real-time notification alerts, and embedded medical attachment previews.")

    # Add AI/LLM Declaration
    doc.add_heading('3. AI/LLM Declaration', level=1)
    doc.add_paragraph("I declare that AI/LLM tools were utilized as an assistive technology during the development of this project. The overall contribution of AI/LLM in code generation, debugging, and drafting boilerplate architectures is estimated at 20%. All generated components were thoroughly tested, refactored, and adapted to align precisely with the project's logic and security requirements.")

    # Frameworks and Libraries
    doc.add_heading('4. Frameworks and Libraries Used', level=1)
    doc.add_paragraph("Backend Technologies:")
    doc.add_paragraph("- Flask (Web Application Framework)\n- Flask-SQLAlchemy (ORM)\n- Flask-Security-Too (Authentication & RBAC)\n- Flask-Migrate (Database Versioning)\n- SQLite DB")
    doc.add_paragraph("Frontend Technologies:")
    doc.add_paragraph("- Vue.js (Component State & Rendering)\n- Bootstrap 5 (CSS framework & UI elements)\n- Chart.js (Interactive visual analytics)\n- jsPDF & html2pdf.js (Dynamic Client-side PDF Generation)")

    # ER Diagram Details
    doc.add_heading('5. ER Diagram Details', level=1)
    doc.add_paragraph("The database revolves around the central `User` model, extending into subsequent profile roles and medical records. Detailed Relationships:")
    doc.add_paragraph("- User (1) to (Many) Role (via roles_users Association Table)")
    doc.add_paragraph("- User (1) to (1) Patient Profile")
    doc.add_paragraph("- User (1) to (1) Doctor Profile")
    doc.add_paragraph("- Department (1) to (Many) Doctors")
    doc.add_paragraph("- Patient (1) to (Many) Appointments")
    doc.add_paragraph("- Doctor (1) to (Many) Appointments")
    doc.add_paragraph("- Appointment (1) to (1) Treatment")
    doc.add_paragraph("- Treatment (1) to (Many) External Attachments")
    
    # API Resource Endpoints
    doc.add_heading('6. Key API Resource Endpoints', level=1)
    doc.add_paragraph("The application operates via several strategic RESTful API endpoints for asynchronous rendering:")
    doc.add_paragraph("- GET /api/admin/dashboard-stats: Aggregates real-time patient mix and workforce distribution data for admin charts.")
    doc.add_paragraph("- GET /api/admin/appointment-stats: Fetches and scopes date-range filtered appointment logs and trend sets.")
    doc.add_paragraph("- GET /patient/api/appointments: Retrieves authenticated patient's historical visit records and current schedule.")
    doc.add_paragraph("- GET /doctor/api/dashboard-data: Provides doctor with a unified payload of their queue and pending requests.")
    doc.add_paragraph("- POST /patient/api/book: Securely reserves a specialized time slot.")
    doc.add_paragraph("- POST /doctor/api/appointment/start/<id>: Triggers clinical initialization of an appointment session.")
    doc.add_paragraph("- POST /doctor/api/appointment/save-session/<id>: Saves structured diagnosis, prescription payload, and medical attachments.")

    # Drive Link
    doc.add_heading('7. Presentation Video Link', level=1)
    doc.add_paragraph("Below is the Google Drive link to the video demonstration:")
    doc.add_paragraph("https://drive.google.com/file/d/1UFNLGKUD4Qwfl0eD304t-r1eoS9v6KI_/view?usp=sharing")

    report_path = os.path.join('/Users/navyasharma/Developer/Projects/hospital-management-version2', 'NavyaSharma_23f1000795.docx')
    doc.save(report_path)
    print("Report named NavyaSharma_23f1000795.docx created successfully.")

if __name__ == '__main__':
    create_report()
